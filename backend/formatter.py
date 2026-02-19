"""
Document formatter: rule-based + optional LLM heading/section detection,
with margins, title page, page breaks, and caption styling.
"""
import json
import os
import re
from collections import Counter
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Optional LLM: set GEMINI_API_KEY or GOOGLE_API_KEY in env to enable AI-based paragraph labeling (Gemini)
# Model can be overridden via GOOGLE_MODEL env var, defaults to gemini-2.5-flash
# Lazy import - only loads when actually needed (when use_llm=True)
_GEMINI_AVAILABLE = False
genai = None
GOOGLE_MODEL = os.environ.get("GOOGLE_MODEL", "gemini-2.5-flash")

def _ensure_gemini():
    """Lazy import of Gemini only when needed."""
    global _GEMINI_AVAILABLE, genai
    if not _GEMINI_AVAILABLE:
        try:
            import google.generativeai as genai
            _GEMINI_AVAILABLE = True
        except ImportError:
            pass
    return _GEMINI_AVAILABLE

# -------- 1. Rule-based heading / caption detection --------

# Numbered sections: 1., 1.1, 1.1.1, 2., etc.
RE_NUMBERED_HEADING = re.compile(r"^\s*\d+(\.\d+)*\.?\s+\S")

# Common section titles (exact or prefix)
SECTION_KEYWORDS = (
    "abstract", "acknowledgement", "acknowledgments", "appendix", "references",
    "bibliography", "contents", "table of contents", "introduction", "conclusion",
    "chapter", "part", "preface", "foreword", "executive summary", "index",
)
RE_CHAPTER = re.compile(r"^\s*(chapter|part)\s+\d+", re.I)
RE_FIGURE_CAPTION = re.compile(r"^\s*(figure|fig\.?)\s*\d*[.:]?\s*\S", re.I)
RE_TABLE_CAPTION = re.compile(r"^\s*(table)\s*\d*[.:]?\s*\S", re.I)


def _is_rule_based_heading(text: str) -> bool:
    if not text or len(text) > 200:
        return False
    t = text.strip()
    # All caps (short enough to be a heading)
    if t.isupper() and len(t) < 120:
        return True
    # Short line (likely title/heading)
    if len(t) < 50:
        # Avoid treating single short sentences as headings
        if t.endswith(".") and t.count(".") >= 1 and not RE_NUMBERED_HEADING.match(t):
            return False
        if RE_NUMBERED_HEADING.match(t):
            return True
        if any(kw in t.lower() for kw in SECTION_KEYWORDS):
            return True
        if RE_CHAPTER.match(t):
            return True
        # Short and no trailing period → likely heading
        if not t.endswith("."):
            return True
    # Numbered section at start
    if RE_NUMBERED_HEADING.match(t):
        return True
    # Known section keyword anywhere at start
    lower = t.lower()
    if any(lower.startswith(kw) or lower.startswith(kw + " ") for kw in SECTION_KEYWORDS):
        return True
    if RE_CHAPTER.match(t):
        return True
    return False


def _is_caption(text: str) -> bool:
    if not text or len(text) > 300:
        return False
    t = text.strip()
    return bool(RE_FIGURE_CAPTION.match(t) or RE_TABLE_CAPTION.match(t))


def _is_likely_title(text: str, is_first: bool) -> bool:
    """First non-empty paragraph that is short and not numbered is often the title."""
    if not text or not is_first:
        return False
    t = text.strip()
    if len(t) > 100 or RE_NUMBERED_HEADING.match(t):
        return False
    if t.isupper() or (len(t) < 80 and not t.endswith(".")):
        return True
    return False


# -------- 2. Optional LLM-based labeling --------

def _get_llm_labels(paragraph_texts: list[str], api_key: str | None) -> list[str] | None:
    """Return list of labels: title, heading, body, caption, other. One per paragraph. Uses Gemini."""
    if not _ensure_gemini() or not api_key or not paragraph_texts:
        return None
    genai.configure(api_key=api_key)
    prompt = """Classify each of the following document paragraphs into exactly one label per line:
- title (document or section title, usually short, one line)
- heading (section heading, subsection title)
- body (normal paragraph)
- caption (figure/table caption)
- other

Return ONLY a JSON array of strings, one label per paragraph, in order. Example: ["title","heading","body","body","caption"]
Paragraphs (one per line, numbered):
"""
    for i, p in enumerate(paragraph_texts, 1):
        prompt += f"{i}. {p[:200]}{'...' if len(p) > 200 else ''}\n"
    try:
        model = genai.GenerativeModel(GOOGLE_MODEL)
        resp = model.generate_content(prompt, generation_config={"temperature": 0})
        content = (resp.text or "").strip()
        content = content.removeprefix("```json").removeprefix("```").strip()
        labels = json.loads(content)
        if isinstance(labels, list) and len(labels) == len(paragraph_texts):
            return [str(x).lower() for x in labels]
    except Exception:
        pass
    return None


def _get_llm_summary(text: str, api_key: str | None, max_chars: int = 4000) -> str | None:
    """Return a 1–2 sentence summary of the document text. Uses Gemini."""
    if not _ensure_gemini() or not api_key or not text.strip():
        return None
    genai.configure(api_key=api_key)
    excerpt = text.strip()[:max_chars]
    if len(text) > max_chars:
        excerpt += "..."
    prompt = f"""Summarize this document in 1–2 short sentences. Be concise.\n\n{excerpt}"""
    try:
        model = genai.GenerativeModel(GOOGLE_MODEL)
        resp = model.generate_content(prompt, generation_config={"temperature": 0.3})
        return (resp.text or "").strip() or None
    except Exception:
        return None


def get_paragraph_labels(input_path: str) -> dict | None:
    """
    Extract non-empty paragraphs from the docx and get Gemini labels only (no formatting).
    Returns {"paragraphs": [{"index": 0, "text_preview": "...", "label": "title"}, ...], "summary": {"title": 1, "heading": 5, ...}}
    or None if Gemini is unavailable or errors.
    """
    doc = Document(input_path)
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    if not texts:
        return {"paragraphs": [], "summary": {}}
    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    labels = _get_llm_labels(texts, api_key)
    if not labels:
        return None
    summary = dict(Counter(labels))
    paragraphs = [
        {"index": i, "text_preview": t[:150] + ("..." if len(t) > 150 else ""), "label": labels[i]}
        for i, t in enumerate(texts)
    ]
    return {"paragraphs": paragraphs, "summary": summary}


def summarize_document(input_path: str, max_chars: int = 4000) -> str | None:
    """Extract text from docx (first max_chars) and return a 1–2 sentence summary from Gemini."""
    doc = Document(input_path)
    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(parts)
    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    return _get_llm_summary(text, api_key, max_chars=max_chars)


# -------- Styling helpers --------

def _apply_run_style(run, size=12, bold=False, italic=False, underline=False):
    run.font.name = "Times New Roman"
    try:
        rPr = run._element.rPr
        if rPr is not None and rPr.rFonts is not None:
            rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    except (AttributeError, TypeError):
        pass
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.underline = underline


def _apply_para_style(para, doc, kind: str, is_after_title: bool = False):
    """Apply style by kind: title, heading, body, caption."""
    # Paragraphs should already have runs (converted before calling this function)
    # But check just in case
    if len(para.runs) == 0:
        # Last resort - try to add a run if somehow we still don't have one
        text = para.text.strip()
        if text:
            para.add_run(text)
        else:
            return  # Empty paragraph - nothing to style
    
    # Do NOT set para.style here - it can clear paragraph content in some docx files.
    # We only set alignment and run formatting, which preserves text.
    
    # Apply base style to all runs first
    for run in para.runs:
        _apply_run_style(run, size=12, bold=False, italic=False, underline=False)

    if kind == "title":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            _apply_run_style(run, size=16, bold=True, italic=False, underline=True)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.line_spacing = 1.0
        return
    if kind == "heading":
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in para.runs:
            _apply_run_style(run, size=14, bold=True, italic=False, underline=True)
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.0
        return
    if kind == "caption":
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            _apply_run_style(run, size=10, bold=False, italic=True, underline=False)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.0
        return
    # body / other
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for run in para.runs:
        _apply_run_style(run, size=12, bold=False, italic=False, underline=False)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.5
    # Page break after title page (next paragraph starts on new page)
    if is_after_title:
        para.paragraph_format.page_break_before = True


def format_document(
    input_path: str,
    output_path: str,
    use_llm: bool = False,
    llm_labels: list[str] | None = None,
) -> None:
    doc = Document(input_path)

    # -------- Page margins (all sections) --------
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # SIMPLIFIED APPROACH: Process paragraphs directly (like old code) but handle runs properly
    # Collect non-empty paragraphs for LLM processing
    paras_with_text: list[tuple] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paras_with_text.append((para, text))

    # Get LLM labels if needed
    if llm_labels is not None and len(llm_labels) == len(paras_with_text):
        pass  # use llm_labels as-is
    elif use_llm:
        api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
        texts = [t for _, t in paras_with_text]
        llm_labels = _get_llm_labels(texts, api_key)
    else:
        llm_labels = None

    # CRITICAL: Convert paragraphs without runs to have runs BEFORE any styling
    # SIMPLEST APPROACH: Just add a run without clearing
    # This might cause text duplication, but GUARANTEES text is preserved
    # para.clear() was causing text loss, so we avoid it entirely
    for para, text in paras_with_text:
        if len(para.runs) == 0:
            # Paragraph has text but no runs - add text as a run
            # Don't clear - just add run (text might duplicate but will be preserved)
            original_text = para.text.strip()
            if original_text:
                para.add_run(original_text)
                # Verify it worked
                if len(para.runs) == 0:
                    # If that failed, something is very wrong
                    raise RuntimeError(f"Failed to add run to paragraph with text: '{original_text[:50]}...'")

    # Process each non-empty paragraph (simpler approach like old code)
    first_para_index = 0
    title_para_index: int | None = None
    
    para_idx = 0  # Index into paras_with_text (non-empty paragraphs only)
    for para, text in paras_with_text:
        # Determine kind
        if llm_labels and para_idx < len(llm_labels):
            kind = llm_labels[para_idx] if llm_labels[para_idx] in ("title", "heading", "body", "caption", "other") else "body"
        else:
            # Rule-based
            if _is_likely_title(text, is_first=(para_idx == first_para_index)):
                kind = "title"
                title_para_index = para_idx
            elif _is_caption(text):
                kind = "caption"
            elif _is_rule_based_heading(text):
                kind = "heading"
            else:
                kind = "body"
        
        # Apply style (paragraphs should now have runs from conversion above)
        is_after_title = title_para_index is not None and para_idx == title_para_index + 1
        _apply_para_style(para, doc, kind, is_after_title=is_after_title)
        
        para_idx += 1

    # Final verification: ensure we didn't lose any text
    total_with_text = len([p for p in doc.paragraphs if p.text.strip()])
    total_paras = len(doc.paragraphs)
    
    print(f"DEBUG: After formatting - Total paragraphs: {total_paras}, With text: {total_with_text}, Processed: {len(paras_with_text)}")
    
    if total_with_text == 0 and len(paras_with_text) > 0:
        # Something went wrong - text was lost
        # Print details about what happened
        print(f"DEBUG: Checking processed paragraphs...")
        for i, (para, text) in enumerate(paras_with_text[:5], 1):
            print(f"  Para {i}: runs={len(para.runs)}, text='{para.text[:50] if para.text else '(empty)'}...'")
        raise ValueError(f"CRITICAL: All text was lost! Had {len(paras_with_text)} paragraphs but document is empty after formatting.")

    doc.save(output_path)
    print(f"DEBUG: Document saved to {output_path}")
