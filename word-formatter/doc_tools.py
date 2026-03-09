"""
Custom browser-use tools for creating and manipulating Word documents (.docx).

These tools allow your Gemini agent to:
- Create formatted Word reports from scratch
- Append sections to existing documents
- Insert images into documents
- Read document text for LLM analysis
"""

import os
from typing import List, Dict

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from browser_use import Controller  # Controller is an alias for Tools in browser-use

controller = Controller()  # This exposes your actions to the Agent


# --- Helper: add a hyperlink run into a paragraph -------------------------------
def _add_hyperlink(paragraph, text: str, url: str):
    """
    Inserts a clickable blue underlined hyperlink run into a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Blue color
    c = OxmlElement("w:color")
    c.set(qn("w:val"), "0000FF")
    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")

    rPr.append(c)
    rPr.append(u)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# --- 1) One-shot "master report" creator ----------------------------------------
@controller.action(
    description=(
        "Create a highly formatted Word report (.docx) with title page, "
        "headers/footers, headings (17pt/13pt), body text, bullets, tables, "
        "hyperlinks, and optional images. "
        "content_blocks is a list of dicts with keys like: "
        "{'type': 'h1'|'h2'|'text'|'bullets'|'table'|'link'|'image', ...}. "
        "Example: [{'type': 'h1', 'text': 'Introduction'}, "
        "{'type': 'table', 'headers': ['Name', 'Price'], 'rows': [['Basic', '$10']]}]"
    )
)
def create_master_doc(
    filename: str,
    title: str,
    content_blocks: List[Dict],
    header_text: str = "Generated Report",
) -> str:
    """
    Creates <filename>.docx from scratch with your full formatting rules.
    
    Args:
        filename: Output filename (without .docx extension)
        title: Title page heading
        content_blocks: List of content dicts (see description for types)
        header_text: Text to put in document header
    
    Returns:
        Success message string
    """
    doc = Document()

    # Global normal style: Times New Roman 12pt
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # Header & footer
    section = doc.sections[0]
    header = section.header
    header.paragraphs[0].text = header_text

    footer = section.footer
    if footer.paragraphs:
        footer.paragraphs[0].text = "Confidential"
    else:
        footer.add_paragraph("Confidential")

    # Title page
    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    for block in content_blocks:
        b_type = block.get("type")

        # Heading 1 (17pt bold)
        if b_type == "h1":
            h = doc.add_heading(block["text"], level=1)
            if h.runs:
                run = h.runs[0]
                run.font.size = Pt(17)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        # Heading 2 (13pt bold)
        elif b_type == "h2":
            h = doc.add_heading(block["text"], level=2)
            if h.runs:
                run = h.runs[0]
                run.font.size = Pt(13)
                run.font.bold = True

        # Body paragraph
        elif b_type == "text":
            p = doc.add_paragraph(block["text"])
            if block.get("bold") and p.runs:
                p.runs[0].bold = True

        # Bulleted list
        elif b_type == "bullets":
            for item in block.get("items", []):
                doc.add_paragraph(item, style="List Bullet")

        # Table: headers + rows
        elif b_type == "table":
            headers = block.get("headers", [])
            rows = block.get("rows", [])
            if headers:
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = "Table Grid"
                hdr_cells = table.rows[0].cells
                for i, h_text in enumerate(headers):
                    hdr_cells[i].text = str(h_text)
                for row_data in rows:
                    row_cells = table.add_row().cells
                    for i, cell_val in enumerate(row_data):
                        row_cells[i].text = str(cell_val)

        # Hyperlink
        elif b_type == "link":
            p = doc.add_paragraph(block.get("pre_text", "Link: "))
            _add_hyperlink(p, block["label"], block["url"])

        # Image from existing file path
        elif b_type == "image":
            path = block.get("path")
            if path and os.path.exists(path):
                width_inches = float(block.get("width_inches", 4.0))
                doc.add_picture(path, width=Inches(width_inches))

    doc.save(f"{filename}.docx")
    return f"Created master document '{filename}.docx'."


# --- 2) Append new sections to an existing report --------------------------------
@controller.action(
    description="Append a new section (heading + body text) to an existing .docx file."
)
def append_section_to_doc(filename: str, heading: str, body: str) -> str:
    """
    Opens <filename>.docx, appends a Heading 2 + body paragraph, and saves.
    
    Args:
        filename: Document filename (without .docx extension)
        heading: Heading 2 text (13pt bold)
        body: Body paragraph text
    
    Returns:
        Success message or error string
    """
    path = f"{filename}.docx"
    if not os.path.exists(path):
        return f"File '{path}' does not exist."

    doc = Document(path)
    h = doc.add_heading(heading, level=2)
    if h.runs:
        run = h.runs[0]
        run.font.size = Pt(13)
        run.font.bold = True
    doc.add_paragraph(body)
    doc.save(path)
    return f"Appended section '{heading}' to {path}."


# --- 3) Insert an image into an existing report ----------------------------------
@controller.action(
    description="Insert an image from a local path into an existing .docx file."
)
def insert_image_into_doc(
    filename: str,
    image_path: str,
    width_inches: float = 4.0,
) -> str:
    """
    Opens <filename>.docx and inserts an image from image_path, if it exists.
    
    Args:
        filename: Document filename (without .docx extension)
        image_path: Full path to image file
        width_inches: Image width in inches (default 4.0)
    
    Returns:
        Success message or error string
    """
    path = f"{filename}.docx"
    if not os.path.exists(path):
        return f"File '{path}' does not exist."
    if not os.path.exists(image_path):
        return f"Image '{image_path}' does not exist."

    doc = Document(path)
    doc.add_picture(image_path, width=Inches(width_inches))
    doc.save(path)
    return f"Inserted image '{image_path}' into {path}."


# --- 4) Read back text from a report (for LLM editing / summarizing) -------------
@controller.action(
    description="Read all text from an existing .docx file and return it as plain text."
)
def get_doc_text(filename: str) -> str:
    """
    Returns the concatenated text of all paragraphs in <filename>.docx.
    
    Args:
        filename: Document filename (without .docx extension)
    
    Returns:
        All document text as plain string, or error message
    """
    path = f"{filename}.docx"
    if not os.path.exists(path):
        return f"File '{path}' does not exist."

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return "\n\n".join(paragraphs)
