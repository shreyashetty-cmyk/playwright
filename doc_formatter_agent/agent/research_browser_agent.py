"""
Vision-based research browser agent using Gemini Flash Vision + Playwright.

High-level behaviour:
- Take a research topic (and desired number of articles).
- Use Playwright to open a search engine and article pages.
- At each step, send a screenshot + structured observation to Gemini Flash Vision.
- Gemini returns a single next action (NAVIGATE, CLICK, SCROLL, EXTRACT_MAIN_TEXT, DONE).
- Execute that action with Playwright until enough articles are collected or DONE.
- Summarise/enhance article texts with text-only Gemini (optional).
- Build a Word report and send it to the existing backend /format endpoint for final formatting.
"""

from __future__ import annotations

import argparse
import os
import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from playwright.sync_api import Page, sync_playwright
import requests

from gemini_client import PlannerAction, call_vision_planner

# Reuse helpers from the existing research agent where appropriate
from research_agent import (  # type: ignore
    _clean_text,
    _extract_main_text,
    _scroll_page_smoothly,
    _setup_console_logging,
    _setup_network_monitoring,
    _summarize_article,
    DEFAULT_NUM_ARTICLES,
    DEMO_SCREENSHOTS_DIR,
    PAGE_TIMEOUT_MS,
    BACKEND_URL,
    FORMAT_ENDPOINT,
)

try:
    from content_enhancer import enhance_content  # type: ignore

    _ENHANCEMENT_AVAILABLE = True
except ImportError:
    enhance_content = None  # type: ignore
    _ENHANCEMENT_AVAILABLE = False


AGENT_DIR = Path(__file__).resolve().parent


def _search_url(topic: str, engine: str) -> str:
    from urllib.parse import quote_plus

    q = quote_plus(topic)
    if engine == "duckduckgo":
        return f"https://duckduckgo.com/?q={q}"
    return f"https://www.bing.com/search?q={q}"


def _summarise_search_results(page: Page, engine: str, max_results: int = 10) -> List[Dict[str, Any]]:
    """
    Extract a lightweight list of search results (title + href) from Bing or DuckDuckGo.
    This is used only to inform Gemini; navigation will use these URLs explicitly.
    """
    results: List[Dict[str, Any]] = []
    try:
        if "bing.com" in (page.url or "") or engine == "bing":
            loc = page.locator("li.b_algo h2 a")
            count = min(max_results, loc.count())
            for i in range(count):
                href = loc.nth(i).get_attribute("href")
                title = (loc.nth(i).inner_text() or "").strip()
                if href and href.startswith("http") and "bing.com" not in href:
                    results.append({"title": title[:200], "url": href})
        else:
            selectors = [
                "article a[data-testid='result-title-a']",
                "a[data-testid='result-title-a']",
                "article a.result__a",
                "a.result__a",
                ".result a",
            ]
            seen = set()
            for selector in selectors:
                loc = page.locator(selector)
                count = min(max_results, loc.count())
                for i in range(count):
                    href = loc.nth(i).get_attribute("href")
                    title = (loc.nth(i).inner_text() or "").strip()
                    if (
                        href
                        and href.startswith("http")
                        and "duckduckgo.com" not in href
                        and href not in seen
                    ):
                        results.append({"title": title[:200], "url": href})
                        seen.add(href)
                    if len(results) >= max_results:
                        break
                if len(results) >= max_results:
                    break
    except Exception:
        pass
    return results


def build_research_observation(
    page: Page,
    topic: str,
    desired_articles: int,
    collected_articles: List[Tuple[str, str]],
    visited_urls: List[str],
    search_engine: str,
) -> Dict[str, Any]:
    """Create an observation dict for Gemini that is tailored to research flows."""
    try:
        url = page.url
    except Exception:
        url = ""
    try:
        title = page.title()
    except Exception:
        title = ""

    try:
        visible_text_prefix = page.inner_text("body")[:2000]
    except Exception:
        visible_text_prefix = ""

    mode = "search" if "bing.com" in url or "duckduckgo.com" in url else "article"
    search_results: List[Dict[str, Any]] = []
    if mode == "search":
        search_results = _summarise_search_results(page, engine=search_engine)

    recent_articles = [
        {"url": u, "chars": len(t)} for (u, t) in collected_articles[-3:]
    ]

    return {
        "topic": topic,
        "desired_articles": desired_articles,
        "collected_articles": len(collected_articles),
        "url": url,
        "title": title,
        "mode": mode,
        "search_engine": search_engine,
        "search_results": search_results,
        "visited_urls": visited_urls[-20:],
        "recent_articles": recent_articles,
        "visible_text_prefix": visible_text_prefix,
    }


def _execute_research_action(
    page: Page,
    action: PlannerAction,
    topic: str,
    articles: List[Tuple[str, str]],
) -> Tuple[bool, Optional[str]]:
    """
    Execute a PlannerAction in the context of research.
    Returns (success, error_message).
    """
    error: Optional[str] = None

    def click_by_description(description: str) -> None:
        text = description
        if ":" in description:
            prefix, rest = description.split(":", 1)
            text = rest.strip()
            prefix = prefix.strip().lower()
            if prefix in {"button", "btn"}:
                page.get_by_role("button", name=text).first.click()
                return
            if prefix in {"link", "a"}:
                page.get_by_role("link", name=text).first.click()
                return
        try:
            page.get_by_text(text, exact=True).first.click()
        except Exception:
            page.get_by_text(text).first.click()

    try:
        if action.action_type == "CLICK":
            if not action.target:
                raise ValueError("CLICK action missing target")
            target = action.target
            if target.startswith("css:"):
                sel = target.split(":", 1)[1]
                page.locator(sel).first.click()
            else:
                click_by_description(target)

        elif action.action_type == "NAVIGATE":
            if not action.target:
                raise ValueError("NAVIGATE action missing URL target")
            page.goto(action.target, wait_until="domcontentloaded", timeout=PAGE_TIMEOUT_MS)

        elif action.action_type == "SCROLL":
            direction = "down"
            pixels = 600
            if action.arguments and isinstance(action.arguments, dict):
                direction = str(action.arguments.get("direction", direction))
                try:
                    pixels = int(action.arguments.get("pixels", pixels))
                except Exception:
                    pixels = 600
            # Use helper for nicer behaviour; invert direction if needed
            if direction.lower().startswith("up"):
                page.mouse.wheel(0, -abs(pixels))
            else:
                page.mouse.wheel(0, abs(pixels))

        elif action.action_type == "EXTRACT_MAIN_TEXT":
            # Let the existing heuristic extractor decide which DOM region is article content.
            text = _extract_main_text(page, scroll_first=False, highlight=False)
            if text.strip():
                url = page.url or ""
                print(f"[research-agent] Extracted article from {url[:80]} ({len(text)} chars)")
                articles.append((url, text))
            else:
                raise ValueError("EXTRACT_MAIN_TEXT produced empty content")

        elif action.action_type == "TYPE":
            if not action.target:
                raise ValueError("TYPE action missing target")
            text = ""
            if action.arguments and isinstance(action.arguments, dict):
                text = str(action.arguments.get("text", ""))
            if not text:
                raise ValueError("TYPE action missing 'text' in arguments")
            target = action.target
            if target.startswith("css:"):
                sel = target.split(":", 1)[1]
                page.fill(sel, text)
            else:
                try:
                    page.get_by_label(target).first.fill(text)
                except Exception:
                    page.locator("input, textarea").first.fill(text)

        elif action.action_type in {"SET_FILE", "DONE"}:
            # SET_FILE is not typically needed for research pages; ignore if produced.
            # DONE is handled by the caller.
            pass

        else:
            raise ValueError(f"Unsupported action_type for research agent: {action.action_type}")

        return True, None
    except Exception as exc:
        error = str(exc)
        print(f"[research-agent] Action {action.action_type} failed: {error}")
        return False, error


def run_research_browser_agent(
    topic: str,
    num_articles: int = DEFAULT_NUM_ARTICLES,
    search_engine: str = "bing",
    headless: bool = False,
    max_steps: int = 40,
) -> Optional[str]:
    """
    Vision-based research flow:
    - Gemini Vision plans navigation + extraction.
    - Text-only Gemini (and content_enhancer) summarize and polish.
    - Document is built and formatted via backend.
    Returns path to formatted docx or None.
    """
    if not topic or not topic.strip():
        print("Error: topic is required.")
        return None

    articles: List[Tuple[str, str]] = []
    visited_urls: List[str] = []

    with sync_playwright() as p:
        launch_options = {
            "headless": headless,
            "slow_mo": 200,
        }
        browser = p.chromium.launch(**launch_options)
        context = browser.new_context(viewport={"width": 1920, "height": 1080})
        page = context.new_page()

        # Optional demo logging (reuse helpers)
        _setup_console_logging(page)
        _setup_network_monitoring(page)

        start_url = _search_url(topic, search_engine)
        print(f"[research-agent] Starting at search URL: {start_url}")
        page.goto(start_url, wait_until="domcontentloaded", timeout=PAGE_TIMEOUT_MS)
        visited_urls.append(start_url)

        for step in range(max_steps):
            if len(articles) >= num_articles:
                print("[research-agent] Reached desired number of articles.")
                break

            screenshot_path = str(AGENT_DIR / f"_research_step_{step:02d}.png")
            page.screenshot(path=screenshot_path, full_page=True)

            obs = build_research_observation(
                page=page,
                topic=topic,
                desired_articles=num_articles,
                collected_articles=articles,
                visited_urls=visited_urls,
                search_engine=search_engine,
            )

            goal = (
                f"Collect up to {num_articles} high-quality articles about the topic: {topic!r}. "
                "Use NAVIGATE with URLs from observation.search_results to open new articles. "
                "When on an article page, use EXTRACT_MAIN_TEXT once per good article. "
                "Use SCROLL if you need to reveal more content. "
                "Stop with DONE when you have collected enough articles or there is nothing useful left."
            )

            print(f"[research-agent] Step {step}: calling Gemini Vision planner...")
            action = call_vision_planner(goal=goal, observation=obs, screenshot_path=screenshot_path)
            print(f"[research-agent]  → action: {action.action_type} target={action.target!r}")

            if action.action_type == "DONE":
                print("[research-agent] Planner reported DONE. Stopping navigation.")
                break

            success, error = _execute_research_action(page, action, topic, articles)

            if page.url and (not visited_urls or page.url != visited_urls[-1]):
                visited_urls.append(page.url)

            if not success:
                # If repeated failures, consider breaking early.
                print(f"[research-agent] Warning: step {step} failed: {error}")

        context.close()
        browser.close()

    if not articles:
        print("[research-agent] No article content could be extracted.")
        return None

    # 2. Summarise/enhance and build Word document (reuse patterns from research_agent).
    print(f"[research-agent] Building document from {len(articles)} article(s)...")
    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    doc = Document()

    doc.add_heading(topic.strip(), level=0)
    if api_key:
        intro = (
            f"This document provides AI-summarized findings from {len(articles)} source(s) "
            f"on the topic: \"{topic.strip()}\"."
        )
    else:
        intro = (
            f"This document summarizes findings from {len(articles)} source(s) "
            f"on the topic: \"{topic.strip()}\"."
        )
    doc.add_paragraph(intro)
    doc.add_paragraph("")

    enhance_content_flag = os.environ.get("ENHANCE_CONTENT", "false").lower() == "true"
    rewrite_style = os.environ.get("REWRITE_STYLE", "academic")

    processed_articles: List[Tuple[str, str]] = []

    for i, (url, raw_text) in enumerate(articles, 1):
        text = raw_text

        # Summarise article with text-only Gemini (optional)
        if api_key:
            print(f"[research-agent] Summarising article {i} with Gemini...")
            summary = _summarize_article(text, api_key)
            if summary:
                text = summary
                print(f"  Article {i}: summarised to {len(text)} chars")

        # Enhance content if enabled
        if enhance_content_flag and _ENHANCEMENT_AVAILABLE and enhance_content:
            enhanced = enhance_content(
                text,
                api_key=api_key,
                rewrite=True,
                rewrite_style=rewrite_style,
                correct_tone_flag=True,
                target_tone="academic",
                check_grammar_flag=True,
            )
            if enhanced.get("enhanced") and enhanced["enhanced"] != text:
                text = enhanced["enhanced"]
                steps_applied = enhanced.get("steps_applied") or []
                if steps_applied:
                    print(f"  ✓ Enhanced article {i}: {', '.join(steps_applied)}")

        processed_articles.append((url, text))

        doc.add_heading(f"Article {i}", level=1)
        paras = [p.strip() for p in text.strip().split("\n\n") if p.strip()]
        for p in paras:
            if len(p) > 30:
                doc.add_paragraph(p)
        doc.add_paragraph(f"Source: {url}")
        doc.add_paragraph("")

    doc.add_heading("References", level=1)
    for i, (url, _) in enumerate(processed_articles, 1):
        doc.add_paragraph(f"{i}. {url}")

    agent_dir = str(AGENT_DIR)
    temp_file = os.path.join(agent_dir, "research_browser_temp.docx")
    doc.save(temp_file)
    print(f"[research-agent] Saved draft to {temp_file}")

    # 3. Send to backend for formatting (reuse approach from research_agent).
    print("[research-agent] Sending document to backend for formatting...")
    backend_proc = None
    try:
        health_check = requests.get(f"{BACKEND_URL}/docs", timeout=5)
        if health_check.status_code != 200:
            raise requests.exceptions.RequestException("Backend not responding")
        print("[research-agent] Backend is running.")
    except requests.exceptions.RequestException:
        print("[research-agent] Backend not running. Attempting to start it...")
        import subprocess

        env = os.environ.copy()
        project_root = os.path.dirname(agent_dir)
        backend_dir = os.path.join(project_root, "backend")

        if not os.path.isdir(backend_dir):
            print(f"Error: Backend directory not found at {backend_dir}")
            out_path = os.path.join(agent_dir, "research_browser_output.docx")
            doc.save(out_path)
            print(f"Unformatted document saved: {out_path}")
            return out_path

        env.setdefault("PYTHONPATH", "")
        if backend_dir not in env["PYTHONPATH"].split(os.pathsep):
            env["PYTHONPATH"] = backend_dir + os.pathsep + env["PYTHONPATH"]

        backend_proc = subprocess.Popen(
            [sys.executable, "-m", "uvicorn", "main:app", "--host", "127.0.0.1", "--port", "8000"],
            cwd=backend_dir,
            env=env,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )

        print("[research-agent] Waiting for backend to start", end="", flush=True)
        import time as _time

        for _ in range(30):
            try:
                if requests.get(f"{BACKEND_URL}/docs", timeout=2).status_code == 200:
                    print(" - ready!")
                    break
            except Exception:
                pass
            print(".", end="", flush=True)
            _time.sleep(1)
        else:
            print("\nFailed to start backend.")
            if backend_proc:
                backend_proc.terminate()
            out_path = os.path.join(agent_dir, "research_browser_output.docx")
            doc.save(out_path)
            print(f"Unformatted document saved: {out_path}")
            return out_path

    try:
        with open(temp_file, "rb") as f:
            files = {
                "file": (
                    "research_browser_temp.docx",
                    f,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            print("[research-agent] Uploading and formatting (this may take 1-2 minutes)...")
            response = requests.post(FORMAT_ENDPOINT, files=files, timeout=300)

        if response.status_code != 200:
            print(f"Backend error: {response.status_code} - {response.text[:200]}")
            raise Exception(f"Backend returned status {response.status_code}")

        out_path = os.path.join(agent_dir, "research_browser_output.docx")
        with open(out_path, "wb") as f:
            f.write(response.content)
        print(f"[research-agent] ✓ Formatted document saved: {out_path}")

        if backend_proc:
            print("[research-agent] Stopping backend...")
            backend_proc.terminate()
            backend_proc.wait(timeout=5)

        return out_path

    except Exception as exc:
        print(f"[research-agent] Error calling backend: {exc}")
        out_path = os.path.join(agent_dir, "research_browser_output_unformatted.docx")
        doc.save(out_path)
        print(f"[research-agent] Unformatted document saved: {out_path}")
        if backend_proc:
            backend_proc.terminate()
        return out_path


def main(argv: Optional[List[str]] = None) -> None:
    parser = argparse.ArgumentParser(
        description="Vision-based research browser agent using Gemini Flash Vision + Playwright."
    )
    parser.add_argument(
        "topic",
        type=str,
        nargs="*",
        help="Research topic (e.g. 'Impact of AI in Intelligence Operations')",
    )
    parser.add_argument(
        "--articles",
        type=int,
        default=DEFAULT_NUM_ARTICLES,
        help=f"Number of articles to collect (default: {DEFAULT_NUM_ARTICLES})",
    )
    parser.add_argument(
        "--duckduckgo",
        action="store_true",
        help="Use DuckDuckGo instead of Bing.",
    )
    parser.add_argument(
        "--headless",
        action="store_true",
        help="Run browser in headless mode.",
    )
    parser.add_argument(
        "--max-steps",
        type=int,
        default=40,
        help="Maximum planner steps before stopping navigation (default: 40).",
    )

    args = parser.parse_args(argv)
    topic = " ".join(args.topic).strip()
    if not topic:
        print("Usage: python research_browser_agent.py --articles N \"Your research topic\"")
        sys.exit(1)

    search_engine = "duckduckgo" if args.duckduckgo else "bing"

    result = run_research_browser_agent(
        topic=topic,
        num_articles=args.articles,
        search_engine=search_engine,
        headless=args.headless,
        max_steps=args.max_steps,
    )
    sys.exit(0 if result else 1)


if __name__ == "__main__":
    main()

