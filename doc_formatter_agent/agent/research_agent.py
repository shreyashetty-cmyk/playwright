"""
Browser Research Agent: given a topic, search the web, extract article text,
summarize using AI (Gemini), build a Word document, and format it using the existing backend.

Usage:
  python research_agent.py "Impact of AI in Intelligence Operations"
  python research_agent.py "Climate change effects on agriculture" --articles 5

Requires: Backend running (or use run_all.py --research "topic")
Requires: GEMINI_API_KEY or GOOGLE_API_KEY environment variable for AI summarization
"""
import os
import sys
import time
from urllib.parse import quote_plus

from docx import Document
from playwright.sync_api import sync_playwright
import requests

# Import semantic memory and content enhancement
try:
    from semantic_memory import get_memory
    from content_enhancer import enhance_content
    _ENHANCEMENT_AVAILABLE = True
except ImportError:
    _ENHANCEMENT_AVAILABLE = False
    get_memory = None
    enhance_content = None

# Optional LLM: set GEMINI_API_KEY or GOOGLE_API_KEY in env to enable AI summarization (Gemini)
# Lazy import - only loads when actually needed
_GEMINI_AVAILABLE = False
genai = None
GOOGLE_MODEL = os.environ.get("GOOGLE_MODEL", "gemini-2.5-flash")

def _ensure_gemini():
    """Lazy import of Gemini only when needed."""
    global _GEMINI_AVAILABLE, genai
    if not _GEMINI_AVAILABLE:
        try:
            import google.generativeai as genai
            _GEMINI_AVAILABLE = True
        except ImportError:
            pass
    return _GEMINI_AVAILABLE

# Backend URL (must be running)
BACKEND_URL = os.environ.get("BACKEND_URL", "http://127.0.0.1:8000")
FORMAT_ENDPOINT = f"{BACKEND_URL}/format"

# Defaults
DEFAULT_NUM_ARTICLES = 3
MAX_CHARS_PER_ARTICLE = 5000
MAX_CHARS_PER_ARTICLE_IN_DOC = 3500  # Trim for neater document
SEARCH_TIMEOUT_MS = 30000  # Increased to 30 seconds for slow networks
PAGE_TIMEOUT_MS = 20000  # Increased to 20 seconds

# Demo features
DEMO_SCREENSHOTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "demo_screenshots")
DEMO_VIDEO_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "demo_videos")


def _clean_text(text: str) -> str:
    """Remove common noise: very short lines, cookie/nav junk, extra whitespace."""
    if not text or not text.strip():
        return ""
    lines = []
    for line in text.splitlines():
        line = line.strip()
        if len(line) < 3:
            continue
        # Skip common UI text
        lower = line.lower()
        if any(s in lower for s in (
            "cookie", "accept all", "privacy policy", "terms of use",
            "subscribe", "newsletter", "sign up", "log in", "menu", "navigation",
            "advertisement", "ads by", "©", "all rights reserved", "click here"
        )):
            continue
        lines.append(line)
    return "\n\n".join(lines) if lines else ""


def _summarize_article(text: str, api_key: str | None, max_input_chars: int = 5000) -> str | None:
    """Summarize article text using Gemini AI. Returns summarized text or None if unavailable."""
    if not _ensure_gemini() or not api_key or not text.strip():
        return None
    try:
        genai.configure(api_key=api_key)
        excerpt = text.strip()[:max_input_chars]
        if len(text) > max_input_chars:
            excerpt += "..."
        prompt = f"""Summarize the following article text in 3-5 concise paragraphs. Focus on key points and main findings. Maintain clarity and coherence.\n\n{excerpt}"""
        model = genai.GenerativeModel(GOOGLE_MODEL)
        resp = model.generate_content(prompt, generation_config={"temperature": 0.3})
        summary = (resp.text or "").strip()
        return summary if summary else None
    except Exception as e:
        print(f"  Warning: AI summarization failed: {e}")
        return None


def _scroll_page_smoothly(page, step_px: int = 300, pause_ms: int = 80, max_scrolls: int = 15) -> None:
    """Scroll the page smoothly so lazy content loads and the flow is visually appealing."""
    try:
        for _ in range(max_scrolls):
            page.mouse.wheel(0, step_px)
            page.wait_for_timeout(pause_ms)
            # Stop if we hit bottom (scroll height doesn't change much)
            try:
                at_bottom = page.evaluate(
                    "() => window.innerHeight + window.scrollY >= document.body.scrollHeight - 10"
                )
                if at_bottom:
                    break
            except Exception:
                pass
        # Scroll back to top so user sees the extracted area
        page.evaluate("window.scrollTo({ top: 0, behavior: 'smooth' })")
        page.wait_for_timeout(400)
    except Exception:
        pass


def _highlight_elements(page, locator, duration_ms: int = 500) -> None:
    """Temporarily highlight elements for visual feedback, then remove after duration_ms."""
    try:
        locator.evaluate_all(
            """(els) => {
            els.forEach(el => {
                el.style.transition = 'background 0.3s ease';
                el.style.backgroundColor = 'rgba(100, 200, 255, 0.25)';
                el.style.outline = '1px solid rgba(100, 200, 255, 0.5)';
            });
        }"""
        )
        page.wait_for_timeout(duration_ms)
        _unhighlight_elements(page, locator)
    except Exception:
        pass


def _highlight_elements_keep(page, locator) -> None:
    """Apply highlight to elements and leave it on (so user sees it while scrolling)."""
    try:
        locator.evaluate_all(
            """(els) => {
            els.forEach(el => {
                el.style.transition = 'background 0.2s ease';
                el.style.backgroundColor = 'rgba(255, 255, 100, 0.35)';
                el.style.outline = '2px solid rgba(255, 180, 0, 0.7)';
            });
        }"""
        )
    except Exception:
        pass


def _unhighlight_elements(page, locator) -> None:
    """Remove highlight from elements."""
    try:
        locator.evaluate_all(
            """(els) => {
            els.forEach(el => {
                el.style.backgroundColor = '';
                el.style.outline = '';
            });
        }"""
        )
    except Exception:
        pass


def _add_progress_overlay(page, message: str, progress: float = 0.0) -> None:
    """Add a visual progress overlay on the page for demo purposes."""
    try:
        page.evaluate(
            f"""
            (() => {{
                let overlay = document.getElementById('playwright-demo-overlay');
                if (!overlay) {{
                    overlay = document.createElement('div');
                    overlay.id = 'playwright-demo-overlay';
                    overlay.style.cssText = `
                        position: fixed;
                        top: 20px;
                        right: 20px;
                        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                        color: white;
                        padding: 15px 25px;
                        border-radius: 10px;
                        box-shadow: 0 8px 32px rgba(0,0,0,0.3);
                        z-index: 999999;
                        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
                        font-size: 14px;
                        max-width: 300px;
                        animation: slideIn 0.3s ease-out;
                    `;
                    document.body.appendChild(overlay);
                }}
                overlay.innerHTML = `
                    <div style="font-weight: 600; margin-bottom: 8px;">🤖 Research Agent</div>
                    <div style="font-size: 13px; opacity: 0.95;">{message}</div>
                    <div style="margin-top: 10px; background: rgba(255,255,255,0.2); height: 4px; border-radius: 2px; overflow: hidden;">
                        <div style="background: white; height: 100%; width: {progress * 100}%; transition: width 0.3s ease;"></div>
                    </div>
                `;
            }})();
            """
        )
    except Exception:
        pass


def _remove_progress_overlay(page) -> None:
    """Remove the progress overlay."""
    try:
        page.evaluate("document.getElementById('playwright-demo-overlay')?.remove()")
    except Exception:
        pass


def _simulate_mouse_movement(page, from_x: int, from_y: int, to_x: int, to_y: int, steps: int = 10) -> None:
    """Simulate smooth mouse movement for demo."""
    try:
        for i in range(steps + 1):
            t = i / steps
            x = int(from_x + (to_x - from_x) * t)
            y = int(from_y + (to_y - from_y) * t)
            page.mouse.move(x, y)
            page.wait_for_timeout(20)
    except Exception:
        pass


def _capture_screenshot(page, filename: str, demo_dir: str = DEMO_SCREENSHOTS_DIR) -> str | None:
    """Capture a screenshot and save it to demo directory."""
    try:
        os.makedirs(demo_dir, exist_ok=True)
        path = os.path.join(demo_dir, filename)
        page.screenshot(path=path, full_page=True)
        return path
    except Exception:
        return None


def _setup_console_logging(page) -> None:
    """Capture and display console logs for demo."""
    def handle_console(msg):
        if msg.type in ("log", "info", "warn", "error"):
            print(f"  [Browser Console] {msg.type.upper()}: {msg.text[:100]}")
    page.on("console", handle_console)


def _setup_network_monitoring(page) -> None:
    """Monitor network requests for demo."""
    requests_count = {"count": 0}
    def handle_request(request):
        requests_count["count"] += 1
        if requests_count["count"] <= 5:  # Show first 5 requests
            print(f"  [Network] {request.method} {request.url[:60]}...")
    page.on("request", handle_request)


def _extract_main_text(page, scroll_first: bool = True, highlight: bool = True) -> str:
    """Extract main content: scroll to load, find content, highlight it, scroll again so user sees it, then extract."""
    try:
        if scroll_first:
            _scroll_page_smoothly(page)
        selectors = [
            "article p",
            "main p",
            "[role='main'] p",
            ".content p",
            ".article-body p",
            ".post-content p",
            "p",
        ]
        for sel in selectors:
            loc = page.locator(sel)
            if loc.count() > 0:
                first_text = "\n".join(loc.all_inner_texts())
                if len(first_text.strip()) > 100:
                    if highlight:
                        _highlight_elements_keep(page, loc)
                    if scroll_first and highlight:
                        _scroll_page_smoothly(page)
                    parts = loc.all_inner_texts()
                    full = "\n".join(parts)
                    if highlight:
                        _unhighlight_elements(page, loc)
                    return _clean_text(full)[:MAX_CHARS_PER_ARTICLE]
    except Exception:
        pass
    return ""


def _search_bing(page, topic: str, num_links: int) -> list[str]:
    """Search Bing and return up to num_links result URLs."""
    query = quote_plus(topic)
    url = f"https://www.bing.com/search?q={query}"
    try:
        # Use networkidle instead of load for faster/better waiting
        page.goto(url, timeout=SEARCH_TIMEOUT_MS, wait_until="domcontentloaded")
        page.wait_for_timeout(2000)
    except Exception as e:
        print(f"  Warning: Bing search page load issue: {e}")
        # Try to continue anyway if page partially loaded
        try:
            page.wait_for_timeout(3000)
        except:
            pass
    
    links = []
    try:
        # Bing organic results
        loc = page.locator("li.b_algo h2 a")
        for i in range(min(num_links, loc.count())):
            href = loc.nth(i).get_attribute("href")
            if href and href.startswith("http") and "bing.com" not in href:
                links.append(href)
    except Exception:
        pass
    return links[:num_links]


def _search_duckduckgo(page, topic: str, num_links: int) -> list[str]:
    """Search DuckDuckGo and return up to num_links result URLs."""
    query = quote_plus(topic)
    url = f"https://duckduckgo.com/?q={query}"
    
    # Retry logic for DuckDuckGo (can be slow or blocked)
    max_retries = 2
    for attempt in range(max_retries):
        try:
            # Use domcontentloaded instead of load - faster and more reliable
            page.goto(url, timeout=SEARCH_TIMEOUT_MS, wait_until="domcontentloaded")
            page.wait_for_timeout(3000)  # Wait for results to render
            break
        except Exception as e:
            if attempt < max_retries - 1:
                print(f"  Retry {attempt + 1}/{max_retries} for DuckDuckGo search...")
                page.wait_for_timeout(2000)
                continue
            else:
                print(f"  Warning: DuckDuckGo search failed after {max_retries} attempts: {e}")
                # Try to continue anyway - page might have partially loaded
                try:
                    page.wait_for_timeout(5000)
                except:
                    pass
    
    links = []
    try:
        # Try multiple selectors for DuckDuckGo (they change frequently)
        selectors = [
            "article a[data-testid='result-title-a']",
            "a[data-testid='result-title-a']",
            "article a.result__a",
            "a.result__a",
            ".result a",
        ]
        
        for selector in selectors:
            try:
                loc = page.locator(selector)
                count = loc.count()
                if count > 0:
                    for i in range(min(num_links, count)):
                        href = loc.nth(i).get_attribute("href")
                        if href and href.startswith("http") and href not in links:
                            links.append(href)
                    if len(links) >= num_links:
                        break
            except Exception:
                continue
        
        # Fallback: try to find any links in results
        if len(links) == 0:
            try:
                all_links = page.locator("a[href^='http']")
                for i in range(min(num_links * 2, all_links.count())):
                    href = all_links.nth(i).get_attribute("href")
                    if href and href.startswith("http") and "duckduckgo.com" not in href and href not in links:
                        links.append(href)
                        if len(links) >= num_links:
                            break
            except Exception:
                pass
    except Exception as e:
        print(f"  Warning: Error extracting DuckDuckGo links: {e}")
    
    return links[:num_links]


def research_topic(
    topic: str,
    num_articles: int = DEFAULT_NUM_ARTICLES,
    search_engine: str = "bing",
    headless: bool = False,
    output_path: str | None = None,
    demo_mode: bool = True,
    record_video: bool = False,
    take_screenshots: bool = True,
) -> str | None:
    """
    Research a topic: search web, extract text from top articles, build docx, format via backend.

    Args:
        topic: Research topic (e.g. "Impact of AI in Intelligence Operations")
        num_articles: Number of articles to open (default 3)
        search_engine: "bing" or "duckduckgo"
        headless: Run browser headless
        output_path: Where to save final formatted docx (default: agent/research_output.docx)
        demo_mode: Enable demo features (progress overlays, visual feedback)
        record_video: Record video of browser automation (saved to demo_videos/)
        take_screenshots: Capture screenshots at key moments (saved to demo_screenshots/)

    Returns:
        Path to saved formatted document, or None on failure.
    """
    if not topic or not topic.strip():
        print("Error: topic is required.")
        return None

    articles_text: list[tuple[str, str]] = []  # (url, text)

    with sync_playwright() as p:
        # Demo mode: slower and with video recording
        launch_options = {
            "headless": headless,
            "slow_mo": 200 if demo_mode else 150,
        }
        if record_video:
            os.makedirs(DEMO_VIDEO_DIR, exist_ok=True)
            video_path = os.path.join(DEMO_VIDEO_DIR, f"research_{topic[:20].replace(' ', '_')}.webm")
            launch_options["record_video_path"] = video_path
        
        browser = p.chromium.launch(**launch_options)
        context = browser.new_context(
            viewport={"width": 1920, "height": 1080},
            record_video_dir=DEMO_VIDEO_DIR if record_video else None,
        )
        page = context.new_page()
        
        # Setup demo features
        if demo_mode:
            _setup_console_logging(page)
            _setup_network_monitoring(page)
            _add_progress_overlay(page, f"🔍 Searching for: {topic}", 0.1)

        # 1. Search
        print(f"Searching for: {topic}")
        if demo_mode:
            _add_progress_overlay(page, f"🔍 Searching {search_engine}...", 0.15)
        
        if search_engine == "duckduckgo":
            links = _search_duckduckgo(page, topic, num_articles)
        else:
            links = _search_bing(page, topic, num_articles)
        
        if demo_mode and take_screenshots:
            _capture_screenshot(page, f"01_search_results_{topic[:20].replace(' ', '_')}.png")
            _add_progress_overlay(page, f"✅ Found {len(links)} articles", 0.3)

        if not links:
            print("No search results found.")
            if demo_mode:
                _remove_progress_overlay(page)
            browser.close()
            return None

        print(f"Found {len(links)} links. Extracting content...")

        # 2. Open each link and extract text (with scroll + highlight)
        for i, url in enumerate(links):
            try:
                if demo_mode:
                    progress = 0.3 + (i / len(links)) * 0.5
                    _add_progress_overlay(page, f"📄 Processing article {i+1}/{len(links)}: {url[:40]}...", progress)
                
                article_page = context.new_page()
                
                if demo_mode:
                    _setup_console_logging(article_page)
                    _add_progress_overlay(article_page, f"🌐 Loading article {i+1}...", 0.1)
                
                article_page.goto(url, timeout=PAGE_TIMEOUT_MS)
                article_page.wait_for_timeout(1800)  # Let page render
                
                if demo_mode:
                    _add_progress_overlay(article_page, f"📖 Extracting content...", 0.5)
                    if take_screenshots:
                        _capture_screenshot(article_page, f"02_article_{i+1}_before_extract.png")
                
                text = _extract_main_text(article_page, scroll_first=True, highlight=True)
                
                if demo_mode:
                    _add_progress_overlay(article_page, f"✅ Content extracted ({len(text)} chars)", 0.8)
                    if take_screenshots:
                        _capture_screenshot(article_page, f"03_article_{i+1}_highlighted.png")
                
                article_page.wait_for_timeout(500)  # Brief pause so user sees highlighted content
                
                if demo_mode:
                    _remove_progress_overlay(article_page)
                
                article_page.close()
                if text.strip():
                    # Store in semantic memory
                    if _ENHANCEMENT_AVAILABLE and get_memory:
                        memory = get_memory()
                        if memory.available:
                            api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
                            # Store original text (will summarize later)
                            memory.store_research(topic, url, text, metadata={"article_index": len(articles_text) + 1})
                            
                            # Cross-verify against previous research
                            verification = memory.cross_verify(text, topic)
                            if verification["verified"]:
                                print(f"  ✓ Cross-verified: {verification['match_count']} similar articles found (confidence: {verification['confidence']:.2f})")
                    
                    # Summarize using AI if available
                    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
                    if api_key:
                        if demo_mode:
                            _add_progress_overlay(page, f"🤖 AI summarizing article {len(articles_text) + 1}...", 0.7 + (i / len(links)) * 0.1)
                        print(f"  Summarizing article {len(articles_text) + 1} with AI...")
                        summarized = _summarize_article(text, api_key)
                        if summarized:
                            text = summarized
                            print(f"  Article {len(articles_text) + 1}: summarized to {len(text)} chars")
                            
                            # Update semantic memory with summary
                            if _ENHANCEMENT_AVAILABLE and get_memory:
                                memory = get_memory()
                                if memory.available:
                                    memory.store_research(topic, url, text, summary=summarized, metadata={"article_index": len(articles_text) + 1, "summarized": True})
                        else:
                            print(f"  Article {len(articles_text) + 1}: AI summary failed, using extracted text ({len(text)} chars)")
                    else:
                        print(f"  Article {len(articles_text) + 1}: {len(text)} chars (no GEMINI_API_KEY, using raw text)")
                    
                    articles_text.append((url, text))
            except Exception as e:
                print(f"  Skip {url[:50]}...: {e}")
                try:
                    article_page.close()
                except Exception:
                    pass

        if demo_mode:
            _add_progress_overlay(page, f"📝 Building document...", 0.85)
            _remove_progress_overlay(page)
        
        context.close()
        browser.close()
        
        if record_video:
            print(f"  📹 Video saved to: {DEMO_VIDEO_DIR}")
        if take_screenshots:
            print(f"  📸 Screenshots saved to: {DEMO_SCREENSHOTS_DIR}")

    if not articles_text:
        print("No article content could be extracted.")
        return None

    # 3. Build Word document (neat structure for formatter)
    doc = Document()
    
    # Title — no empty paragraph after, so no gap below title
    doc.add_heading(topic.strip(), level=0)
    # Introduction (directly under title)
    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if api_key:
        intro = f"This document provides AI-summarized findings from {len(articles_text)} source(s) on the topic: \"{topic.strip()}\"."
    else:
        intro = f"This document summarizes findings from {len(articles_text)} source(s) on the topic: \"{topic.strip()}\"."
    doc.add_paragraph(intro)
    doc.add_paragraph("")

    # Content enhancement options
    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    enhance_content_flag = os.environ.get("ENHANCE_CONTENT", "false").lower() == "true"
    rewrite_style = os.environ.get("REWRITE_STYLE", "academic")
    
    for i, (url, text) in enumerate(articles_text, 1):
        # Article heading
        doc.add_heading(f"Article {i}", level=1)
        
        # Enhance content if enabled
        if enhance_content_flag and _ENHANCEMENT_AVAILABLE and enhance_content:
            enhanced = enhance_content(
                text,
                api_key=api_key,
                rewrite=True,
                rewrite_style=rewrite_style,
                correct_tone_flag=True,
                target_tone="academic",
                check_grammar_flag=True
            )
            if enhanced.get("enhanced") and enhanced["enhanced"] != text:
                text = enhanced["enhanced"]
                if enhanced.get("steps_applied"):
                    print(f"  ✓ Enhanced article {i}: {', '.join(enhanced['steps_applied'])}")
        
        # Text is already summarized if AI was available, so use it directly
        # Split into paragraphs for a neater layout (double newline = new para)
        paras = [p.strip() for p in text.strip().split("\n\n") if p.strip()]
        for p in paras:
            # Skip very short fragments (single line leftovers)
            if len(p) > 30:
                doc.add_paragraph(p)
        
        # Source on its own line (formatter will treat as body)
        doc.add_paragraph(f"Source: {url}")
        doc.add_paragraph("")

    # References section
    doc.add_heading("References", level=1)
    for i, (url, _) in enumerate(articles_text, 1):
        doc.add_paragraph(f"{i}. {url}")

    agent_dir = os.path.dirname(os.path.abspath(__file__))
    temp_file = os.path.join(agent_dir, "research_temp.docx")
    doc.save(temp_file)
    print(f"Saved draft to {temp_file}")

    # 4. Send to backend for formatting
    print("Sending document to backend for formatting...")
    
    # Check if backend is running, start if needed
    backend_proc = None
    try:
        health_check = requests.get(f"{BACKEND_URL}/docs", timeout=5)
        if health_check.status_code != 200:
            raise requests.exceptions.RequestException("Backend not responding")
        print("Backend is running.")
    except requests.exceptions.RequestException:
        print("Backend not running. Attempting to start it...")
        import subprocess
        import sys as sys_module
        env = os.environ.copy()
        # Get backend directory: go up from agent/ to project root, then to backend/
        project_root = os.path.dirname(agent_dir)
        backend_dir = os.path.join(project_root, "backend")
        
        # Verify backend directory exists
        if not os.path.isdir(backend_dir):
            print(f"Error: Backend directory not found at {backend_dir}")
            print("Please ensure backend/ directory exists in the project root.")
            out_path = output_path or os.path.join(agent_dir, "research_output.docx")
            doc.save(out_path)
            print(f"Unformatted document saved: {out_path}")
            print("To format it later:")
            print("  1. Start backend manually: cd backend && uvicorn main:app --reload")
            print(f"  2. Use: python format_document.py {temp_file}")
            return out_path
        
        env.setdefault("PYTHONPATH", "")
        if backend_dir not in env["PYTHONPATH"].split(os.pathsep):
            env["PYTHONPATH"] = backend_dir + os.pathsep + env["PYTHONPATH"]
        
        backend_proc = subprocess.Popen(
            [sys_module.executable, "-m", "uvicorn", "main:app", "--host", "127.0.0.1", "--port", "8000"],
            cwd=backend_dir,
            env=env,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        
        # Wait for backend to be ready
        print("Waiting for backend to start", end="", flush=True)
        for _ in range(30):
            try:
                if requests.get(f"{BACKEND_URL}/docs", timeout=2).status_code == 200:
                    print(" - ready!")
                    break
            except:
                pass
            print(".", end="", flush=True)
            time.sleep(1)
        else:
            print("\nFailed to start backend.")
            if backend_proc:
                backend_proc.terminate()
            out_path = output_path or os.path.join(agent_dir, "research_output.docx")
            doc.save(out_path)
            print(f"Unformatted document saved: {out_path}")
            print("To format it later:")
            print("  1. Start backend: cd backend && uvicorn main:app --reload")
            print(f"  2. Use: python format_document.py {temp_file}")
            return out_path
    
    # Now format the document
    try:
        with open(temp_file, "rb") as f:
            files = {"file": ("research_temp.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            # Increased timeout for large documents and LLM processing
            print("Uploading and formatting (this may take 1-2 minutes)...")
            response = requests.post(FORMAT_ENDPOINT, files=files, timeout=300)  # 5 minutes
        
        if response.status_code != 200:
            print(f"Backend error: {response.status_code} - {response.text[:200]}")
            raise Exception(f"Backend returned status {response.status_code}")
        
        # 5. Save formatted file
        out_path = output_path or os.path.join(agent_dir, "research_output.docx")
        with open(out_path, "wb") as f:
            f.write(response.content)
        print(f"✓ Formatted document saved: {out_path}")
        
        # Clean up backend if we started it
        if backend_proc:
            print("Stopping backend...")
            backend_proc.terminate()
            backend_proc.wait(timeout=5)
        
        return out_path
        
    except requests.exceptions.Timeout:
        print("Warning: Backend request timed out (document may be large).")
        print("Saving unformatted document. You can format it later.")
        out_path = output_path or os.path.join(agent_dir, "research_output.docx")
        doc.save(out_path)
        print(f"Unformatted document saved: {out_path}")
        print("To format it later:")
        print("  1. Start backend: cd backend && uvicorn main:app --reload")
        print(f"  2. Use: python format_document.py {temp_file}")
        if backend_proc:
            backend_proc.terminate()
        return out_path
    except requests.exceptions.ConnectionError:
        print("Error: Backend connection failed. Saving unformatted document.")
        out_path = output_path or os.path.join(agent_dir, "research_output.docx")
        doc.save(out_path)
        print(f"Unformatted document saved: {out_path}")
        print("To format it later:")
        print("  1. Start backend: cd backend && uvicorn main:app --reload")
        print(f"  2. Use: python format_document.py {temp_file}")
        if backend_proc:
            backend_proc.terminate()
        return out_path
    except Exception as e:
        print(f"Error calling backend: {e}")
        print("Saving unformatted document as fallback...")
        out_path = output_path or os.path.join(agent_dir, "research_output_unformatted.docx")
        doc.save(out_path)
        print(f"Unformatted document saved: {out_path}")
        if backend_proc:
            backend_proc.terminate()
        return out_path

    # Optional: remove temp
    try:
        os.remove(temp_file)
    except Exception:
        pass

    return out_path


def main():
    """CLI entry point."""
    args = sys.argv[1:]
    if not args:
        print("Usage: python research_agent.py \"Your research topic\" [options]")
        print("\nOptions:")
        print("  --articles N          Number of articles (1-10, default: 3)")
        print("  --duckduckgo          Use DuckDuckGo instead of Bing")
        print("  --headless            Run browser in headless mode")
        print("  --no-demo             Disable demo features (progress overlays, etc.)")
        print("  --record-video        Record video of browser automation")
        print("  --no-screenshots     Don't capture screenshots")
        sys.exit(1)

    topic_parts = []
    num_articles = DEFAULT_NUM_ARTICLES
    search_engine = "bing"
    headless = False
    demo_mode = True
    record_video = False
    take_screenshots = True

    i = 0
    while i < len(args):
        a = args[i]
        if a == "--articles" and i + 1 < len(args):
            num_articles = max(1, min(10, int(args[i + 1])))
            i += 2
            continue
        if a == "--duckduckgo":
            search_engine = "duckduckgo"
            i += 1
            continue
        if a == "--headless":
            headless = True
            i += 1
            continue
        if a == "--no-demo":
            demo_mode = False
            i += 1
            continue
        if a == "--record-video":
            record_video = True
            i += 1
            continue
        if a == "--no-screenshots":
            take_screenshots = False
            i += 1
            continue
        topic_parts.append(a)
        i += 1

    topic = " ".join(topic_parts).strip()
    if not topic:
        print("Error: topic is required.")
        sys.exit(1)

    result = research_topic(
        topic,
        num_articles=num_articles,
        search_engine=search_engine,
        headless=headless,
        demo_mode=demo_mode,
        record_video=record_video,
        take_screenshots=take_screenshots,
    )
    sys.exit(0 if result else 1)


if __name__ == "__main__":
    main()
